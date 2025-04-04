import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import io
from datetime import datetime
import math
import re
from docx.shared import Inches, Pt


#def rerun():
    #raise st.script_runner.RerunException(st.script_request_queue.RerunData(None))
# Function to add customer details
if "rerun_counter" not in st.session_state:
    st.session_state.rerun_counter = 0

# Function to trigger rerun by incrementing a counter
def rerun():
    st.session_state.rerun_counter += 1

def set_page_size(doc, width_mm=148, height_mm=210):
    # Convert mm to twips (1 mm = 567 twips)
    width_twips = int(width_mm * 567)
    height_twips = int(height_mm * 567)

    # Access the document's section properties
    for section in doc.sections:
        # Get the section's XML element
        sect_pr = section._sectPr

        # Create or modify the pgSize element
        pg_size = OxmlElement('w:pgSize')
        pg_size.set(qn('w:w'), str(width_twips))
        pg_size.set(qn('w:h'), str(height_twips))

        # Remove any existing pgSize element
        for existing_pg_size in sect_pr.findall(qn('w:pgSize')):
            sect_pr.remove(existing_pg_size)

        # Add the new pgSize element
        sect_pr.append(pg_size)

def add_customer_details(doc, customer_name, address, phone, product_name,bill_number,Total,add_notes,Sub_Total,Total_amount,Shipping_Cost):




    table = doc.add_table(rows=2, cols=2)
    left_cell = table.cell(0, 0)
    left_paragraph = left_cell.add_paragraph( "Order Number:"+bill_number)
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_paragraph.runs[0]
    left_run.font.size = Pt(12)
    table1 = doc.add_table(rows=1, cols=3)
    # Add customer name
    cell1 = table.cell(1, 0)
    cell1.text = "Customer Name:"+customer_name
    cell1_run = cell1.paragraphs[0].runs[0]
    cell1_run.font.size = Pt(12)
    # Add address
    cell3 = table1.cell(0, 0)
    cell3.text = "To Address:\n"+address
    cell3_run = cell3.paragraphs[0].runs[0]
    cell3_run.font.size = Pt(12)
    # Add address
    cell4 = table1.cell(0, 2)
    shop_add="Krish Accessories\nGG Nagar, Nerkundram\nChennai 600107\n8939789237"
    cell4.text = "From Address:\n"+shop_add
    cell4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell4_run = cell4.paragraphs[0].runs[0]
    cell4_run.font.size = Pt(12)
    table2 = doc.add_table(rows=2, cols=1)
    # Add phone
    cell5 = table2.cell(0, 0)
    cell5.text = "Customer Phone:"+phone
    cell5_run = cell5.paragraphs[0].runs[0]
    cell5_run.font.size = Pt(12)


    cell6_1 = table2.cell(1, 0)
    cell6_1.text = "•Unboxing video is mandatory for returns or replacements."
    cell6_1_run = cell6_1.paragraphs[0].runs[0]
    cell6_1_run.font.size = Pt(12)


    data = [["S.No", "Product Name","QTY", "Unit_Total"]]
    product_list = product_name.split(",")

    total_list=Sub_Total.split(",")
    total_sum_list = [float(x) for x in Sub_Total.split(",")]
    total_sum = sum(total_sum_list)
    for idx, product in enumerate(product_list, start=1):
        match = re.search(r'\[(\d+)\]', product)
        number_in_brackets = match.group(1)
        product_full_name = re.sub(r'\s*\[\d+\]', '', product).strip()
        total_value = total_list[idx-1]
        data.append([idx, product_full_name, number_in_brackets, total_value ])

    table_grid = doc.add_table(rows=1, cols=len(data[0]))
    table_grid.style = 'Table Grid'  # Apply a grid style to the table
    # Disable AutoFit to manually control column widths
    # Set the column widths explicitly
    table_grid.columns[0].width = Inches(0.5)  # S.No column
    table_grid.columns[1].width = Inches(2.5)  # Product Name column
    table_grid.columns[2].width = Inches(1.0)  # QTY column
    table_grid.columns[3].width = Inches(1.5)  # Total column




    # Add Headers to the Table
    hdr_cells = table_grid.rows[0].cells
    for i, header in enumerate(data[0]):
        table_grid.autofit = False
        hdr_cells[i].text = str(header)
        hdr_paragraph = hdr_cells[i].paragraphs[0]
        hdr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add Table Rows
# Add table rows
    for row_data in data[1:]:
        row = table_grid.add_row().cells
        table_grid.autofit = False
        for i, value in enumerate(row_data):
            row[i].text = str(value)
            row_paragraph = row[i].paragraphs[0]
            #row_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    table_cost = doc.add_table(rows=2, cols=3)
    cell_ship_cost = table_cost.cell(0, 2)
    cell_ship_cost.text = "Shipping Cost:"+str(Shipping_Cost)
    cell_ship_cost_run = cell_ship_cost.paragraphs[0].runs[0]
    cell_ship_cost_run.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell_ship_cost_run.font.size = Pt(12)
    cell_cost = table_cost.cell(1, 2)
    cell_cost.text = "Total Cost:"+str(Total_amount)
    cell_cost_run = cell_cost.paragraphs[0].runs[0]
    cell_cost_run.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell_cost_run.font.size = Pt(12)





    if add_notes and not isinstance(add_notes, float) or (isinstance(add_notes, float) and not math.isnan(add_notes)):
        table_notes = doc.add_table(rows=1, cols=1)
        notes_cell = table_notes.cell(0,0)
        notes_paragraph = notes_cell.add_paragraph(str("Customer Notes:"+add_notes))
        #grand_total_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        grand_notes_paragraph = notes_paragraph.runs[0]
        grand_notes_paragraph.font.size = Pt(12)

    byte_stream = BytesIO()
    doc.save(byte_stream)  # Save the document into the BytesIO stream
    byte_stream.seek(0)
    return byte_stream
# Insert page break before the next customer if needed
    if i < len(Prepaid_order) - 1:
        doc.add_page_break()

def convert_df_to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')
    processed_data = output.getvalue()
    return processed_data

# Function to generate PDF from DataFrame
def dataframe_to_pdf(df):
    buffer = io.BytesIO()  # Create an in-memory bytes buffer

    # Create a figure and axis for the table
    fig, ax = plt.subplots(figsize=(12, 14))  # Adjust size to fit the content
    ax.axis('off')  # Hide the axis

    # Render the DataFrame as a table
    table = ax.table(
        cellText=df.values,
        colLabels=df.columns,
        cellLoc='center',
        loc='center',
    )

    # Adjust table properties
    table.auto_set_font_size(False)
    table.set_fontsize(12)  # Adjust font size
    table.auto_set_column_width(col=list(range(len(df.columns))))  # Adjust column width

    # Save to PDF
    with PdfPages(buffer) as pdf:
        pdf.savefig(fig, bbox_inches='tight')  # Save the current figure

    plt.close(fig)  # Close the figure to free memory
    buffer.seek(0)  # Reset buffer pointer to the beginning
    return buffer

# Title of the app
st.title("Excel File Uploader")

# Header for single file upload
st.header('Single File Upload')

# Initialize session state for download
if 'downloaded' not in st.session_state:
    st.session_state.downloaded = False

# File uploader for Excel files
data = st.file_uploader("Upload your Excel file here:", type=['xlsx'])

if data is not None and not st.session_state.downloaded:
    # Read the uploaded Excel file
    try:
        order = pd.read_excel(data, sheet_name='Orders')
        inquiry = pd.read_excel(data, sheet_name='Inquiries')
        st.success("File uploaded and read successfully!")

        true_order = order[
    (order['Confirmed Order'].astype(str).str.lower() == 'true') &
    (order['Order Status'].astype(str).str.upper() == 'COMPLETED')
]
        true_order["Payment Mode"] = true_order["Payment Mode"].astype(str).str.upper()
        true_order.loc[true_order["Payment Mode"].str.upper() == "RAZORPAY", "Payment Method"] = "Prepaid"
        true_order.loc[true_order["Payment Mode"].str.upper() == "PARTIAL COD", "Payment Method"] = "PCOD"
        true_order.loc[true_order["Payment Mode"].astype(str).str.upper() == "CASH ON DELIVERY", "Payment Method"] = "COD"
        true_inquiry = inquiry[
    (inquiry['Confirmed Order'].astype(str).str.lower() == 'true') &
    (inquiry['Order Status'].astype(str).str.upper() == 'COMPLETED')
]
        true_inquiry ['Qty Ordered'] = inquiry['Product Name'] + " "+ "["+ inquiry['Item Count'].astype(str) + "]"
        true_order_number = true_order['Order Number'].astype(str).str.upper()
        filtered_inquiry_df = true_inquiry[true_inquiry['Order Number'].astype(str).str.upper().isin(true_order_number)]

        merged_inquiry_df = filtered_inquiry_df.groupby('Order Number', as_index=False).agg({
    'Product Name': lambda x: ', '.join(x.astype(str)),
    'Price': lambda x: ', '.join(x.astype(str)),
    'Qty Ordered': lambda x: ','.join(x.astype(str)),
    'Product Weight': lambda x: x.sum() * 1000  # convert weight to grams
})

        column_names = [
            '*Sale Order Number', '*Pickup Location Name', '*Transport Mode',
            '*Payment Mode', 'COD Amount', '*Customer Name', '*Customer Phone',
            '*Shipping Address Line1', 'Shipping Address Line2', '*Shipping City',
            '*Shipping State', '*Shipping Pincode', '*Item Sku Code',
            '*Item Sku Name', '*Quantity Ordered', 'Packaging Type',
            '*Unit Item Price', 'Length (cm)', 'Breadth (cm)', 'Height (cm)',
            'Weight (gm)', 'Fragile Shipment', 'Discount Type', 'Discount Value',
            'Tax Class Code', 'Customer Email',
            'Billing Address same as Shipping Address', 'Billing Address Line1',
            'Billing Address Line2', 'Billing City', 'Billing State',
            'Billing Pincode', 'e-Way Bill Number', 'Seller Name',
            'Seller GST Number', 'Seller Address Line1', 'Seller Address Line2',
            'Seller City', 'Seller State', 'Seller Pincode','Shipping Cost'
        ]
        cust_column_names = ['Tdate','CNo','DeptDesc','CrCode','RefNo','cnee','Caddr1',
        'Caddr2','Caddr3','CPincode','CPhone','Destn','Wt','Pcs','DDate','Status','RName','NonDStatus']
        dispatch_column_names = ['Tdate','CrCode','RefNo','cnee','CPincode']
        upload = pd.DataFrame(columns=column_names)
        cust_upload = pd.DataFrame(columns=cust_column_names)
        dispatch_upload = pd.DataFrame(columns=dispatch_column_names)
        new_rows = []
        cust_rows =[]
        dispatch_rows=[]

        for index, row in merged_inquiry_df.iterrows():
            group_order = order[order['Order Number'] == row['Order Number']]
        #    payment_mode = true_order.loc[true_order["Order Number"] == row['Order Number'], ["Payment Method","Additional Notes"]].values[0]
        #    payment_mode = true_order.loc[true_order["Order Number"] == row['Order Number'], ["Payment Method", "Additional Notes"]].values[0]
            payment_mode = true_order.loc[true_order["Order Number"] == row["Order Number"], ["Payment Method", "Additional Notes","Shipping Cost","Total Amount"]].values[0]


            if not group_order.empty:
                customer_phone = str(group_order['Customer Mobile Number'].values[0]).replace('91', '', 1).strip()
                cod_amount = None
                if (payment_mode =="PCOD").any():
                    cod_amount= int(group_order['Total Amount'].values[0]) - 400
                elif ((payment_mode =="COD").any()):
                    cod_amount= int(group_order['Total Amount'].values[0])









                # Create a new row dictionary
                new_row = {
                    '*Sale Order Number': row['Order Number'],
                    '*Pickup Location Name': 'KRISH ACCESSORIES D2C',
                    '*Transport Mode': "Surface",
                    '*Payment Mode':  payment_mode[0],
                    'COD Amount': cod_amount,
                    '*Customer Name': group_order['Customer Name'].values[0],
                    '*Customer Phone': customer_phone,
                    '*Shipping Address Line1': group_order['Shipping Address'].values[0],
                    'Shipping Address Line2': "",
                    '*Shipping City': group_order['City'].values[0],
                    '*Shipping State': group_order['State'].values[0],
                    '*Shipping Pincode': group_order['Pincode'].values[0],
                    '*Item Sku Code': row['Product Name'],
                    '*Item Sku Name': row['Qty Ordered'],
                    '*Quantity Ordered': "1",
                    'Packaging Type': "",
                    'Sub_Total': row['Price'],
                    '*Unit Item Price': int(float(group_order['Total Amount'].values[0])),
                    'Length (cm)': "10",
                    'Breadth (cm)': "10",
                    'Height (cm)': "10",
                    'Weight (gm)': row['Product Weight'],
                    'Fragile Shipment': "",
                    'Discount Type': "",
                    'Discount Value': "",
                    'Tax Class Code': "",
                    'Customer Email': "",
                    'Billing Address same as Shipping Address': "",
                    'Billing Address Line1': "",
                    'Billing Address Line2': "",
                    'Billing City': "",
                    'Billing State': "",
                    'Billing Pincode': "",
                    'e-Way Bill Number': "",
                    'Seller Name': "",
                    'Seller GST Number': "",
                    'Seller Address Line1': "",
                    'Seller Address Line2': "",
                    'Seller City': "",
                    'Seller State': "",
                    'Seller Pincode': "",
                    'Notes':payment_mode[1],
                    'Shipping Cost': int(float(payment_mode[2])),
                    'Total Cost':int(float(payment_mode[3]))
                }

                word_row ={
                "SL" : "",
                "Product Name" :row['Product Name'],
                "Qty" :row['Qty Ordered'],
                "Total":group_order['Total Amount'].values[0],
                "Notes":group_order['Additional Notes'].values[0]
                }

                # Create a new cust row dictionary
                cust_row = {
                    'Tdate':datetime.now().strftime("%Y-%m-%d"),
                    'CNo':"",
                    'DeptDesc':"",
                    'CrCode':"MVY0254",
                    'RefNo': row['Order Number'],
                    'cnee':group_order['Customer Name'].values[0],
                    'Caddr1':"",
                    'Caddr2':"",
                    'Caddr3':group_order['State'].values[0],
                    'CPincode':group_order['Pincode'].values[0],
                    'CPhone':"",
                    'Destn':group_order['State'].values[0],
                    'Wt':row['Product Weight'],
                    'Pcs':"1",
                    'DDate':"",
                    'Status':"",
                    'RName':"",
                    'NonDStatus':""

                }
                dispatch_row= {
                    'Tdate':datetime.now().strftime("%Y-%m-%d"),
                    'CrCode':"MVY0254",
                    'RefNo': row['Order Number'],
                    'cnee':group_order['Customer Name'].values[0],
                    'CPincode':group_order['Pincode'].values[0]
                }


                # Append the new row to new_rows list
                new_rows.append(new_row)
                cust_rows.append(cust_row)
                dispatch_rows.append(dispatch_row)

        upload = pd.concat([upload, pd.DataFrame(new_rows)], ignore_index=True)
        cust_upload = pd.concat([cust_upload, pd.DataFrame(cust_rows)], ignore_index=True)
        dispatch_upload = pd.concat([dispatch_upload, pd.DataFrame(dispatch_rows)], ignore_index=True)
        # Generate PDF for the updated DataFrame
        pdf_buffer = dataframe_to_pdf(dispatch_upload)



        doc = Document()
        set_page_size(doc)
        payment_mode = upload.loc[upload["*Payment Mode"].str.upper() == "PREPAID"]
        Prepaid_order=payment_mode[["*Customer Name","*Shipping Address Line1","*Customer Phone","*Item Sku Name","*Sale Order Number","*Unit Item Price","Notes",'Sub_Total','Shipping Cost','Total Cost']]
        Prepaid_order = Prepaid_order.rename(columns={
    '*Customer Name': 'Customer_Name',
    '*Sale Order Number':'Order_Number',
    '*Shipping Address Line1': 'Address',
    '*Customer Phone': 'Customer_Phone',
    '*Item Sku Name': 'Product_Name',
    '*Unit Item Price':'Total',
    'Notes':'Notes',
    'Sub_Total':'Sub_Total',
    'Total_Amount':'Total Cost',
    'Shipping Cost':'Shipping Cost'
})

        for i in range(len(Prepaid_order)):
            row = Prepaid_order.iloc[i]
   # print(row)

    # Correctly reference the columns without extra spaces
            customer_name = row['Customer_Name']
            address = row['Address']  # Make sure column names match exactly
            customer_phone = row['Customer_Phone']
            product_name = row['Product_Name']
            order_number=row['Order_Number']
            Unit_Total =row['Total']
            Notes=row['Notes']
            Sub_Total=row['Sub_Total']
            Total=row['Total Cost']
            Shipping_Cost=row['Shipping Cost']


    # Call the function to add customer details to the document
            word_data=add_customer_details(doc, customer_name, address, customer_phone, product_name,order_number,int(Unit_Total,Notes),int(Sub_Total),int(Total),int(Shipping_Cost))
            #if (i + 1) % 2 == 0:
            doc.add_page_break()




        # Display the first few rows of the DataFrame
        #st.write(upload.head())

        # Provide a download button
        excel_data = convert_df_to_excel(upload)
        #excel_data = upload.to_excel('upload.xlsx', index=False)
        cust_excel_data = convert_df_to_excel(cust_upload)

        if st.download_button(
                    label="Download data as Excel",
                    data=excel_data,
                    file_name='Output.xlsx',
                    #mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                ):
                    st.session_state.downloaded = True
                    rerun()
        if st.download_button(
                    label="Download booking data as Excel",
                    data=cust_excel_data,
                    file_name='booking.xlsx',
                    #mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                ):
                    st.session_state.downloaded = True
                    rerun()

        if st.download_button(
                    label="Download  booking data as PDF",
                    data=pdf_buffer,
                    file_name="booking_data.pdf",
                    mime="application/pdf"
                ):
                    st.session_state.downloaded = True
                    rerun()





        if st.download_button(
            label="Download data as word",
            data=word_data,
            file_name='Output.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        ):
            st.session_state.downloaded = True
            rerun()

    except Exception as e:
        st.error(f"Error reading the Excel file: {e}")

# Reset the state for a new upload
if st.session_state.downloaded:
    st.session_state.downloaded = False
    rerun()
